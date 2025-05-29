import streamlit as st
from docx import Document
from datetime import date

st.set_page_config(page_title="HydroTerra Product Evaluation", layout="centered")

st.title("🔍 HydroTerra Product Evaluation Tool")

# Step 1 – Product Info
st.header("1. Enter Product Info")
product_name = st.text_input("Product Name")
model_title = st.text_input("Model or Page Title")
product_url = st.text_input("Product URL")
summary = st.text_area("Product Summary (1–2 sentences)")

# Step 2 – Generate Prompt
if product_name and model_title and product_url and summary:
    st.header("2. Copy Prompt for ChatGPT")
    prompt = f"""
You are an environmental technology analyst for HydroTerra.

Evaluate the following product and provide a structured report with:
1. Product Overview (including service gap, innovation, and market alignment)
2. Supplier and Distribution Considerations
3. Competitive Landscape (whether the supplier competes with Solinst or Aquaread)
4. Product Portfolio Review (a table with product name, description, HydroTerra fit, and a rating out of 5)
5. Overall Recommendation and Next Steps
6. A scoring summary out of 35 points:
   - Service Gap (10)
   - Innovation (5)
   - Market Fit (10)
   - Supplier Accessibility (5)
   - Integration Potential (5)

Product Name: {product_name}  
Title: {model_title}  
URL: {product_url}  
Summary: {summary}
"""
    st.code(prompt, language="markdown")
    st.info("👉 Copy this prompt into ChatGPT (free version) and paste the response below.")

    # Step 3 – Paste GPT Output
    st.header("3. Paste GPT Output")
    gpt_response = st.text_area("Paste the full GPT-generated report here")

    # Step 4 – Generate Word Report
    if gpt_response:
        st.header("4. Export Report")
        if st.button("Generate Word Report"):
            doc = Document()
            doc.add_heading(f"{product_name} – Product Evaluation", 0)
            doc.add_paragraph(f"Generated: {date.today().strftime('%d %B %Y')}")
            doc.add_paragraph(f"Product URL: {product_url}")
            doc.add_paragraph("")

            for line in gpt_response.split("\n"):
                if line.strip().startswith("1.") or line.strip().startswith("2.") or \
                   line.strip().startswith("3.") or line.strip().startswith("4.") or \
                   line.strip().startswith("5.") or line.strip().startswith("6."):
                    doc.add_heading(line.strip(), level=1)
                else:
                    doc.add_paragraph(line.strip())

            filename = f"{product_name.replace(' ', '_')}_Product_Evaluation_HydroTerra.docx"
            doc.save(filename)
            with open(filename, "rb") as file:
                st.download_button(
                    label="📄 Download Word Report",
                    data=file,
                    file_name=filename,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
